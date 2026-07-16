"""Module chứa các hàm tiện ích cho pipeline tạo hợp đồng.

Các hàm:
    - ask_llm_with_template: Gọi vLLM với câu hỏi user + template content.
    - save_response_to_word: Lưu response LLM thành file Word (.docx) với định dạng chuẩn.
"""

import logging
import os
import re

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

logger = logging.getLogger(__name__)

# --- Hằng số cấu hình ---
LLM_TEMPERATURE     = 0.1     # Gần deterministic — LLM faithfully copy template, ít sáng tạo
LLM_MAX_TOKENS      = int(os.getenv("LLM_MAX_TOKENS", "128000"))  # Giới hạn output token tối đa
LLM_TIMEOUT         = float(os.getenv("LLM_TIMEOUT", "3600.0"))    # Timeout giây
LLM_ENABLE_THINKING = False   # Qwen3: False = nhanh hơn, tiết kiệm VRAM
LLM_CONTEXT_WINDOW  = int(os.getenv("LLM_CONTEXT_WINDOW", "128000"))  # Context window tổng
_CHARS_PER_TOKEN    = 2.5     # Qwen3 tokenizer: tiếng Việt ~2.5 ký tự/token

# Font & style mặc định cho hợp đồng tiếng Việt
DEFAULT_FONT_NAME = "Times New Roman"
DEFAULT_FONT_SIZE = Pt(13)
HEADING1_FONT_SIZE = Pt(16)
HEADING2_FONT_SIZE = Pt(14)
HEADING3_FONT_SIZE = Pt(13)
CONTRACT_OUTPUT_DIR = os.getenv("FOLDER_PATH_CONTRACT", "database/storage/contract")




# =============================================================================
# BƯỚC 1: PROMPT LLM TỐI ƯU
# =============================================================================


SYSTEM_PROMPT = """\
Bạn là chuyên gia soạn thảo hợp đồng pháp lý tại Việt Nam. Nhiệm vụ của bạn là TÁI TẠO NGUYÊN VẸN mẫu hợp đồng và điền thông tin theo yêu cầu.

QUY TẮC TUYỆT ĐỐI — VI PHẠM LÀ SAI:
1. CHỈ xuất nội dung hợp đồng thuần túy. KHÔNG lời chào, KHÔNG giải thích, KHÔNG ghi chú.
2. SAO CHÉP NGUYÊN VẸN 100% mẫu hợp đồng — từng chữ, từng câu, từng điều khoản, từng dấu câu.
3. CHỈ thay đổi phần thông tin người dùng yêu cầu điền. Thông tin nào thiếu → giữ "..." hoặc "___".
4. TUYỆT ĐỐI KHÔNG tóm tắt, rút gọn, paraphrase, hay bỏ qua bất kỳ nội dung nào dù chỉ 1 câu.
5. PHẢI sinh đủ toàn bộ hợp đồng từ đầu đến điều khoản ký kết cuối cùng — không được dừng giữa chừng.

ĐỊNH DẠNG MARKDOWN:
- Quốc hiệu/tên hợp đồng: ## heading 2
- Tiêu đề điều khoản: ## heading 2 | Mục con: ### heading 3
- In đậm **text**: tên bên, chức danh, số tiền, ngày tháng quan trọng
- Danh sách: dùng - hoặc 1. 2. 3.
- Bảng: cú pháp | col1 | col2 | --- KHÔNG dùng HTML, KHÔNG dùng code block.
"""


async def ask_llm_with_template(user_input: str, template_content: str, model_name: str = None) -> str:
    """
    Gọi vLLM với câu hỏi user + template content sử dụng streaming để tránh timeout.

    Input:
        user_input (str): Câu hỏi / yêu cầu của người dùng.
        template_content (str): Nội dung markdown của template hợp đồng.

    Output:
        str: Nội dung hợp đồng markdown từ LLM.

    Raises:
        RuntimeError: Nếu vLLM API trả response rỗng hoặc không hợp lệ.
    """

    # Đặt YÊU CẦU lên trước để LLM nắm task trước khi đọc template
    # Thêm instruction nhắc lại ở cuối để reinforcement
    user_content = (
        f"YÊU CẦU: {user_input}\n\n"
        f"MẪU HỢP ĐỒNG CẦN TÁI TẠO (sao chép nguyên văn, chỉ điền thông tin theo yêu cầu trên):\n\n"
        f"{template_content}\n\n"
        f"--- NHẮC LẠI ---\n"
        f"Hãy tái tạo TOÀN BỘ hợp đồng trên từ đầu đến cuối, nguyên văn từng điều khoản, "
        f"chỉ thay thế thông tin theo yêu cầu: {user_input}"
    )

    # Tắt thinking mode Qwen3 khi không cần (nhanh hơn ~2-3x, tiết kiệm VRAM)
    extra_body = {}
    if not LLM_ENABLE_THINKING:
        extra_body["chat_template_kwargs"] = {"enable_thinking": False}

    # --- Dùng streaming để tránh timeout với hợp đồng siêu dài ---
    chunks = []
    finish_reason = None
    prompt_tokens = completion_tokens = total_tokens = 0

    # Load prompt hệ thống từ DB theo feature (DB-only)
    from service.runtime_config_service import (
        get_required_active_prompt_content,
        resolve_model_runtime,
        PROMPT_FEATURE_CONTRACT_TEMPLATE,
    )

    current_system_prompt = await get_required_active_prompt_content(
        PROMPT_FEATURE_CONTRACT_TEMPLATE
    )

    # --- Tính max_tokens động để tránh vượt context window ---
    # Tính tổng ký tự của TOÀN BỘ prompt gửi lên (system + user)
    total_prompt_chars = len(current_system_prompt) + len(user_content)
    # Ước tính input tokens: Qwen3 ~2.5 ký tự/token cho tiếng Việt, thêm safety buffer 3000
    estimated_input_tokens = int(total_prompt_chars / _CHARS_PER_TOKEN) + 3000
    available_output_tokens = LLM_CONTEXT_WINDOW - estimated_input_tokens
    dynamic_max_tokens = max(256, min(LLM_MAX_TOKENS, available_output_tokens))

    logger.info(
        "Prompt: %d ký tự ≈ %d tokens → max_tokens=%d (context=%d)",
        total_prompt_chars, estimated_input_tokens, dynamic_max_tokens, LLM_CONTEXT_WINDOW,
    )

    client, resolved_model, meta = await resolve_model_runtime(model_name)
    stream = await client.chat.completions.create(
        model=resolved_model,
        messages=[
            {"role": "system", "content": current_system_prompt},
            {"role": "user", "content": user_content},
        ],
        temperature=LLM_TEMPERATURE,
        max_tokens=dynamic_max_tokens,
        timeout=LLM_TIMEOUT,
        extra_body=extra_body or None,
        stream=True,
        stream_options={"include_usage": True},
    )

    async for chunk in stream:
        # Nhận token usage từ chunk cuối
        if hasattr(chunk, "usage") and chunk.usage:
            u = chunk.usage
            prompt_tokens = u.prompt_tokens
            completion_tokens = u.completion_tokens
            total_tokens = u.total_tokens

        if not chunk.choices:
            continue

        choice = chunk.choices[0]

        # Lấy finish_reason từ chunk cuối
        if choice.finish_reason:
            finish_reason = choice.finish_reason

        delta = choice.delta
        if delta and delta.content:
            token = delta.content
            chunks.append(token)

    # Log token usage
    if total_tokens:
        logger.info(
            "Token usage — prompt: %d, completion: %d, total: %d",
            prompt_tokens, completion_tokens, total_tokens,
        )

    # Cảnh báo nếu LLM bị cắt ngắn
    if finish_reason == "length":
        logger.warning(
            "⚠️  finish_reason=length — hợp đồng bị cắt ngắn! "
            "completion_tokens=%d/%d. Cần tăng LLM_MAX_TOKENS hoặc rút ngắn template.",
            completion_tokens, LLM_MAX_TOKENS,
        )
    else:
        logger.info("finish_reason=%s — hợp đồng sinh hoàn chỉnh ✓", finish_reason)

    response = "".join(chunks).strip()

    if not response:
        raise RuntimeError("vLLM API trả response rỗng hoặc không hợp lệ")

    # Loại bỏ think block nếu có (Qwen3 thinking mode)
    response = re.sub(r"<think>.*?</think>", "", response, flags=re.DOTALL).strip()

    return response


# =============================================================================
# BƯỚC 1.5: PROMPT LLM CHO AI FAST (KHÔNG CÓ TEMPLATE)
# =============================================================================

FAST_SYSTEM_PROMPT = """\
Bạn là Luật sư Trưởng chuyên nghiệp tại Việt Nam. Nhiệm vụ của bạn là soạn thảo một hợp đồng pháp lý HOÀN CHỈNH từ con số Không dựa trên yêu cầu ngắn gọn của người dùng.

QUY TẮC TUYỆT ĐỐI — VI PHẠM LÀ SAI:
1. CHỈ xuất nội dung hợp đồng thuần túy. KHÔNG lời chào, KHÔNG giải thích, KHÔNG ghi chú. Chữ đầu tiên sinh ra phải bắt đầu bằng CỘNG HÒA XÃ HỘI CHỦ NGHĨA...
2. Hợp đồng phải ĐẦY ĐỦ các phần: Quốc hiệu/Tiêu ngữ, Tên hợp đồng, Thông tin các bên, Các điều khoản trọng yếu (Giá cả, Thanh toán, Quyền/Nghĩa vụ, Phạt vi phạm, Căn cứ pháp lý, Ký kết).
3. Thông tin nào thiếu trong yêu cầu → hãy để trống dạng "..." hoặc "[...]". TUYỆT ĐỐI làm rõ các phần đã được User nhắc tới.
4. PHẢI sinh đủ từ đầu đến điều khoản ký kết cuối cùng (có chỗ ký tên rành mạch).

ĐỊNH DẠNG MARKDOWN (TUYỆT ĐỐI TUÂN THỦ):
- Quốc hiệu/tên hợp đồng: ## heading 2
- Tiêu đề Điều khoản: ## heading 2 | Mục con: ### heading 3
- In đậm **text**: tên bên, chức danh, số tiền, ngày tháng quan trọng
- Bảng liệt kê: dùng Markdown table | Cột 1 | Cột 2 |
"""

async def ask_llm_fast(user_input: str, model_name: str = None) -> str:
    """
    Gọi vLLM để phác thảo hợp đồng nhanh không cần template.
    """
    user_content = f"YÊU CẦU LẬP HỢP ĐỒNG:\n{user_input}\n\nHãy soạn thảo hợp đồng thật chi tiết và đầy đủ."

    extra_body = {}
    if not LLM_ENABLE_THINKING:
        extra_body["chat_template_kwargs"] = {"enable_thinking": False}

    chunks = []
    finish_reason = None
    prompt_tokens = completion_tokens = total_tokens = 0

    from service.runtime_config_service import (
        get_required_active_prompt_content,
        resolve_model_runtime,
        PROMPT_FEATURE_CONTRACT_FAST,
    )

    current_fast_prompt = await get_required_active_prompt_content(
        PROMPT_FEATURE_CONTRACT_FAST
    )

    total_prompt_chars = len(current_fast_prompt) + len(user_content)
    estimated_input_tokens = int(total_prompt_chars / _CHARS_PER_TOKEN) + 500
    available_output_tokens = LLM_CONTEXT_WINDOW - estimated_input_tokens
    dynamic_max_tokens = max(1000, min(8000, available_output_tokens)) # Fast output cap

    logger.info("Fast Prompt: %d ký tự → max_tokens=%d", total_prompt_chars, dynamic_max_tokens)

    client, resolved_model, meta = await resolve_model_runtime(model_name)
    stream = await client.chat.completions.create(
        model=resolved_model,
        messages=[
            {"role": "system", "content": current_fast_prompt},
            {"role": "user", "content": user_content},
        ],
        temperature=0.7, # Cho phép AI sáng tạo một chút
        max_tokens=dynamic_max_tokens,
        timeout=LLM_TIMEOUT,
        extra_body=extra_body or None,
        stream=True,
        stream_options={"include_usage": True},
    )

    async for chunk in stream:
        if hasattr(chunk, "usage") and chunk.usage:
            u = chunk.usage
            prompt_tokens, completion_tokens, total_tokens = u.prompt_tokens, u.completion_tokens, u.total_tokens

        if not chunk.choices: continue
        choice = chunk.choices[0]
        if choice.finish_reason: finish_reason = choice.finish_reason
        delta = choice.delta
        if delta and delta.content: chunks.append(delta.content)

    if total_tokens:
        logger.info("Fast AI Token usage: prompt %d, completion %d, total %d", prompt_tokens, completion_tokens, total_tokens)

    response = "".join(chunks).strip()
    if not response:
        raise RuntimeError("vLLM Fast API trả response rỗng")

    response = re.sub(r"<think>.*?</think>", "", response, flags=re.DOTALL).strip()
    return response


# System prompt ngắn gọn cho toàn tắt hợp đồng
_SUMMARY_SYSTEM_PROMPT = (
    "Bạn là trợ lý hợp đồng. Hãy tóm tắt ngắn gọn hợp đồng dưới đây trong 3-5 câu, "
    "bao gồm: tên loại hợp đồng, các bên tham gia, giá trị/thời hạn chính, và nội dung cốt lõi. "
    "Viết bằng tiếng Việt, súc tích, dễ hiểu."
)


async def stream_contract_summary(
    contract_text: str,
    sse_queue,          # asyncio.Queue
    user_id: str = "",
    session_id: int = -1,
    model_name: str = None,
) -> str:
    """
    Gọi LLM để tóm tắt hợp đồng ngắn gọn và push từng token qua sse_queue.
    Trả về toàn bộ nội dung tóm tắt.
    """
    logger.info("stream_contract_summary: bắt đầu tóm tắt hợp đồng")

    # Giời hạn đầu vào để tránh vượt context
    max_input_chars = 40000
    truncated = contract_text[:max_input_chars]
    if len(contract_text) > max_input_chars:
        truncated += "\n...[nội dung quá dài, đã cắt bớt]"

    extra_body = {}
    if not LLM_ENABLE_THINKING:
        extra_body["chat_template_kwargs"] = {"enable_thinking": False}

    from service.runtime_config_service import (
        get_required_active_prompt_content,
        resolve_model_runtime,
        PROMPT_FEATURE_CONTRACT_SUMMARY,
    )

    current_summary_prompt = await get_required_active_prompt_content(
        PROMPT_FEATURE_CONTRACT_SUMMARY
    )

    full_summary = ""
    client, resolved_model, meta = await resolve_model_runtime(model_name)
    try:
        stream = await client.chat.completions.create(
            model=resolved_model,
            messages=[
                {"role": "system", "content": current_summary_prompt},
                {"role": "user", "content": truncated},
            ],
            temperature=0.3,
            max_tokens=512,   # Tóm tắt ngắn — 500 token là đủ
            timeout=120,
            extra_body=extra_body or None,
            stream=True,
        )

        async for chunk in stream:
            if not chunk.choices:
                continue
            delta = chunk.choices[0].delta
            if delta and delta.content:
                token = delta.content
                full_summary += token
                if sse_queue:
                    await sse_queue.put({
                        "user_id": user_id,
                        "session_id": session_id,
                        "title": "Đang tóm tắt...",
                        "mess": token,
                        "end": False,
                    })

        logger.info("stream_contract_summary: xong, %d ký tự", len(full_summary))

    except Exception as e:
        logger.error("stream_contract_summary lỗi: %s", e)
        full_summary = f"(Không thể tóm tắt: {e})"
        if sse_queue:
            await sse_queue.put({
                "user_id": user_id,
                "session_id": session_id,
                "title": "Đang tóm tắt...",
                "mess": full_summary,
                "end": False,
            })

    if sse_queue:
        await sse_queue.put(None)

    return full_summary



# BƯỚC 2: TIỀN XỬ LÝ MARKDOWN (SỬA HEADING BỊ VỠ TỪ DOCLING/LLM)
# =============================================================================

# Regex heading: ## hoặc ### theo sau bởi text
_HEADING_RE = re.compile(r"^(#{1,3})\s+(.*)")

# Pattern nhận diện heading hợp lệ trong hợp đồng VN
_VALID_HEADING_PATTERNS = re.compile(
    r"(?i)"
    r"(CỘNG\s*HÒA|Độc\s*lập|HỢP\s*ĐỒNG|PHẦN\s*\d|Điều\s*\d|Mục\s*\d"
    r"|BÊN\s*[AB]|Chủ\s*đầu\s*tư|Nhà\s*thầu|Một\s*bên|Bên\s*kia"
    r"|HAI\s*BÊN|THANH\s*TOÁN|TRÁCH\s*NHIỆM|THỎA\s*THUẬN"
    r"|NỘI\s*DUNG|GIÁ\s*TRỊ|THỜI\s*GIAN|VI\s*PHẠM|TẠM\s*DỪNG"
    r"|CHẤM\s*DỨT|ĐIỀU\s*KHOẢN|HỒ\s*SƠ|CĂN\s*CỨ"
    r"|CAM\s*KẾT|QUYỀN|NGHĨA\s*VỤ|BẢO\s*HÀNH|BẢO\s*HIỂM"
    r"|PHẠT|BỒI\s*THƯỜNG|GIẢI\s*QUYẾT|TRANH\s*CHẤP"
    r"|PHỤ\s*LỤC|KÝ\s*KẾT|ĐẠI\s*DIỆN)"
)

# Độ dài tối thiểu cho nội dung heading (sau khi bỏ ##)
_MIN_HEADING_CONTENT_LEN = 4


def _is_valid_heading(content: str) -> bool:
    """Kiểm tra nội dung sau ## có phải heading hợp lệ hay chỉ là fragment bị vỡ."""

    content = content.strip()

    # Quá ngắn → fragment
    if len(content) < _MIN_HEADING_CONTENT_LEN:
        return False

    # Bắt đầu bằng chữ thường → rất có thể là tiếp nối (VD: "sơ bàn giao", "a các bên")
    if content and content[0].islower():
        return False

    # Khớp pattern heading hợp đồng → chắc chắn hợp lệ
    if _VALID_HEADING_PATTERNS.search(content):
        return True

    # Viết hoa toàn bộ + đủ dài → nhiều khả năng là heading (VD: "CÔNG TY ABC", "ĐẠI DIỆN BÊN A")
    alpha_chars = re.sub(r"[^a-zA-ZÀ-ỹ]", "", content)
    if len(alpha_chars) >= 5 and content == content.upper():
        return True

    # Bắt đầu bằng số + dấu chấm → heading đánh số (VD: "1. Chủ đầu tư", "2. Nhà thầu")
    if re.match(r"^\d+\.\s+", content):
        return True

    # Mặc định: nếu đủ dài (>= 15 ký tự) và bắt đầu bằng chữ hoa → heading
    if len(content) >= 15 and content[0].isupper():
        return True

    # Còn lại → coi là fragment
    return False


def _preprocess_markdown(text: str) -> str:
    """
    Tiền xử lý markdown trước khi chuyển sang Word.
    Sửa các vấn đề từ Docling:
        1. Merge heading bị vỡ thành 2 dòng liên tiếp.
        2. Demote heading giả (fragment quá ngắn hoặc chữ thường).
        3. Dọn dẹp artifact: <!-- image -->, - - (double dash).
    """

    # Bước 0: Dọn artifact
    text = re.sub(r"<!--.*?-->", "", text)  # Xóa HTML comments
    text = re.sub(r"^- -", "-", text, flags=re.MULTILINE)  # "- -text" → "- text"

    lines = text.split("\n")
    result = []
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        heading_match = _HEADING_RE.match(stripped)
        if heading_match:
            prefix = heading_match.group(1)  # "#", "##", or "###"
            content = heading_match.group(2).strip()

            # Peek dòng tiếp theo: nếu cũng là heading cùng level và nội dung ngắn/fragment → merge
            while i + 1 < len(lines):
                next_stripped = lines[i + 1].strip()
                next_match = _HEADING_RE.match(next_stripped)

                if next_match and next_match.group(1) == prefix:
                    next_content = next_match.group(2).strip()
                    # Fragment → merge vào dòng hiện tại
                    if not _is_valid_heading(next_content):
                        content = content + " " + next_content
                        i += 1
                        continue
                break

            # Validate heading sau khi merge
            if _is_valid_heading(content):
                result.append(f"{prefix} {content}")
            else:
                # Demote thành text thường (giữ content, bỏ ##)
                result.append(content)
        else:
            result.append(line)

        i += 1

    return "\n".join(result)


# =============================================================================
# BƯỚC 3 + 4: WORD CONVERTER VỚI ĐỊNH DẠNG ĐẦY ĐỦ
# =============================================================================


def _setup_document() -> Document:
    """Tạo Document với page setup chuẩn hợp đồng VN (A4, margins, font mặc định)."""

    doc = Document()

    # Page setup: A4, margins chuẩn hợp đồng VN
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    # Font mặc định
    style = doc.styles["Normal"]
    font = style.font
    font.name = DEFAULT_FONT_NAME
    font.size = DEFAULT_FONT_SIZE
    font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.space_after = Pt(3)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.5   # Chuẩn hợp đồng VN: 1.5 dòng

    # Style cho headings
    for level, size, bold in [
        ("Heading 1", HEADING1_FONT_SIZE, True),
        ("Heading 2", HEADING2_FONT_SIZE, True),
        ("Heading 3", HEADING3_FONT_SIZE, True),
    ]:
        h_style = doc.styles[level]
        h_style.font.name = DEFAULT_FONT_NAME
        h_style.font.size = size
        h_style.font.bold = bold
        h_style.font.color.rgb = RGBColor(0, 0, 0)
        h_style.paragraph_format.space_before = Pt(6)
        h_style.paragraph_format.space_after = Pt(3)

    return doc


# Regex để tách inline formatting: **bold** và *italic*
_INLINE_PATTERN = re.compile(
    r"(\*\*(.+?)\*\*)"   # group 1,2: bold
    r"|(\*(.+?)\*)"      # group 3,4: italic
)


def _add_formatted_runs(paragraph, text: str):
    """Thêm runs vào paragraph với hỗ trợ **bold** và *italic*.
    Font name/size kế thừa từ paragraph style — không set lại từng run.
    """

    last_end = 0
    for match in _INLINE_PATTERN.finditer(text):
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])

        if match.group(2):      # **bold**
            paragraph.add_run(match.group(2)).bold = True
        elif match.group(4):    # *italic*
            paragraph.add_run(match.group(4)).italic = True

        last_end = match.end()

    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def _add_heading(doc: Document, text: str, level: int):
    """Thêm heading với font và alignment chuẩn."""

    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.name = DEFAULT_FONT_NAME
    run.font.color.rgb = RGBColor(0, 0, 0)

    if level <= 2:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return heading


def _add_table(doc: Document, table_lines: list[str]):
    """Parse markdown table lines và thêm vào document."""

    rows_data = []
    for line in table_lines:
        line = line.strip().strip("|")
        # Bỏ dòng separator (|---|---|)
        if re.match(r"^[\s\-:|]+$", line):
            continue
        cells = [cell.strip() for cell in line.split("|")]
        rows_data.append(cells)

    if not rows_data:
        return

    num_cols = max(len(row) for row in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.style = "Table Grid"

    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = ""
                para = cell.paragraphs[0]
                _add_formatted_runs(para, cell_text)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.space_before = Pt(0)

    # Bold header row
    if rows_data:
        for cell in table.rows[0].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True


def save_response_to_word(llm_response: str, contract_name: str) -> dict:
    """
    Lưu response LLM (markdown) thành file Word (.docx) với định dạng chuẩn.

    Hỗ trợ: headings, **bold**, *italic*, bullet list, numbered list, table, dòng kẻ.

    Input:
        llm_response (str): Nội dung markdown từ LLM.
        contract_name (str): Tên file hợp đồng (không đuôi).

    Output:
        dict: {"status": str, "path_name": str}
    """

    output_filename = f"{contract_name}.docx"
    os.makedirs(CONTRACT_OUTPUT_DIR, exist_ok=True)
    output_path = os.path.join(CONTRACT_OUTPUT_DIR, output_filename)

    # Tiền xử lý: sửa heading bị vỡ, merge fragment, dọn artifact
    llm_response = _preprocess_markdown(llm_response)

    doc = _setup_document()

    lines = llm_response.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # --- Dòng trống: bỏ qua để tránh giãn cách quá mức ---
        # (space_after trên paragraph đã tạo khoảng cách đủ)
        if not stripped:
            i += 1
            continue

        # --- Markdown table (bắt đầu bằng |) ---
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _add_table(doc, table_lines)
            continue

        # --- Dòng kẻ ngang --- hoặc ___
        if re.match(r"^(-{3,}|_{3,}|\*{3,})$", stripped):
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            run = para.add_run("─" * 60)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(128, 128, 128)
            i += 1
            continue

        # --- Heading ### ---
        if stripped.startswith("### "):
            _add_heading(doc, stripped[4:], level=3)
            i += 1
            continue

        # --- Heading ## ---
        if stripped.startswith("## "):
            _add_heading(doc, stripped[3:], level=2)
            i += 1
            continue

        # --- Heading # ---
        if stripped.startswith("# "):
            _add_heading(doc, stripped[2:], level=1)
            i += 1
            continue

        # --- Numbered list (1. 2. 3.) ---
        num_match = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if num_match:
            para = doc.add_paragraph(style="List Number")
            _add_formatted_runs(para, num_match.group(2))
            i += 1
            continue

        # --- Bullet list (- hoặc *) ---
        bullet_match = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet_match:
            para = doc.add_paragraph(style="List Bullet")
            _add_formatted_runs(para, bullet_match.group(1))
            i += 1
            continue

        # --- Paragraph thường: thụt đầu dòng chuẩn hợp đồng VN ---
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Pt(28)  # ~1cm thụt đầu dòng
        _add_formatted_runs(para, stripped)
        i += 1

    doc.save(output_path)
    logger.info("Đã xuất file hợp đồng: %s", output_filename)

    return {
        "status": "ok",
        "path_name": output_filename,
    }
